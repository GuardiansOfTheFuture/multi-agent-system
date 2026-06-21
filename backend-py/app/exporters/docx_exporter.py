from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from .base import Exporter


class DocxExporter(Exporter):
    def export(self, markdown: str, title: str) -> bytes:
        doc = Document()
        doc.add_heading(title, 0)
        for line in markdown.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                p = doc.add_paragraph(line.strip()[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", line.strip()):
                doc.add_paragraph(line.strip(), style="List Number")
            elif line.strip():
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                doc.add_paragraph(clean)
            else:
                doc.add_paragraph()
        import io
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
